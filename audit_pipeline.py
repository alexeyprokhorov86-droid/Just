#!/usr/bin/env python3
"""
audit_pipeline.py — ночной ревизор пайплайна.

Проверяет и исправляет:
1. TG-вложения без media_analysis / content_text (берёт из S3, fallback Telegram)
2. Email-вложения без анализа
3. source_documents без distillation
4. source_documents без source_chunks

Шлёт отчёт в Telegram.

Cron: 0 2 * * * cd /home/admin/telegram_logger_bot && .../python audit_pipeline.py >> .../audit_pipeline.log 2>&1
"""

import os
import sys
import io
import time
import asyncio
import logging
import base64
import hashlib
import fcntl
import psycopg2
import psycopg2.extras
import requests
import subprocess
from proxy_config import get_proxy_dict

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

from dotenv import load_dotenv
load_dotenv('/home/admin/telegram_logger_bot/.env')

BOT_TOKEN = os.getenv('BOT_TOKEN', '')
ROUTERAI_API_KEY = os.getenv('ROUTERAI_API_KEY', '')
ROUTERAI_BASE_URL = os.getenv('ROUTERAI_BASE_URL', 'https://routerai.ru/api/v1')
ADMIN_USER_ID = os.getenv('ADMIN_USER_ID', '')
BUCKET_NAME = os.getenv('ATTACHMENTS_BUCKET_NAME', '')

DB_CONFIG = {
    'host': os.getenv('DB_HOST', '172.20.0.2'),
    'port': 5432,
    'dbname': 'knowledge_base',
    'user': 'knowledge',
    'password': os.getenv('DB_PASSWORD'),
}

EXCLUDED_TABLES = {'tg_chat_1003360907471_torty_otgruzki'}
VENV_PYTHON = '/home/admin/telegram_logger_bot/venv/bin/python'
SCRIPTS_DIR = '/home/admin/telegram_logger_bot'

MAX_MEDIA_FIX = 30
MAX_DISTILL_BATCHES = 50
MAX_CHUNKS_DOCS = 500


def get_conn():
    return psycopg2.connect(**DB_CONFIG)


def get_s3_client():
    try:
        import boto3
        from botocore.config import Config as BotoConfig
        return boto3.client(
            service_name='s3',
            endpoint_url=os.getenv('ATTACHMENTS_BUCKET_ENDPOINT', 'https://s3.cloud.ru'),
            region_name=os.getenv('ATTACHMENTS_BUCKET_REGION', 'ru-central-1'),
            aws_access_key_id=os.getenv('ATTACHMENTS_BUCKET_ACCESS_KEY', ''),
            aws_secret_access_key=os.getenv('ATTACHMENTS_BUCKET_SECRET_KEY', ''),
            config=BotoConfig(s3={"addressing_style": "path"})
        )
    except Exception as e:
        logger.warning(f"S3 client error: {e}")
        return None


def get_gpt_client():
    from openai import OpenAI
    return OpenAI(api_key=ROUTERAI_API_KEY, base_url=ROUTERAI_BASE_URL, timeout=300)


# ============================================================
# СКАЧИВАНИЕ ФАЙЛОВ: S3 -> Telegram fallback
# ============================================================

async def download_file(item, s3, bot):
    """Скачать файл: сначала S3, потом Telegram. Вернуть (bytes, filename, source)."""
    
    # 1. Попробовать S3
    if item.get('storage_path') and s3:
        try:
            path = item['storage_path']
            if path.startswith('s3://'):
                parts = path[5:].split('/', 1)
                bucket = parts[0]
                key = parts[1]
                resp = s3.get_object(Bucket=bucket, Key=key)
                data = resp['Body'].read()
                filename = key.split('/')[-1]
                # Убираем хэш-префикс из имени
                if '_' in filename and len(filename.split('_')[0]) == 16:
                    filename = '_'.join(filename.split('_')[1:])
                return data, filename, 's3'
        except Exception as e:
            logger.warning(f"S3 download failed: {e}")
    
    # 2. Fallback на Telegram
    if item.get('file_id') and bot:
        try:
            f = await bot.get_file(item['file_id'])
            data = await f.download_as_bytearray()
            filename = f.file_path.split('/')[-1] if f.file_path else f"file_{item['message_id']}"
            
            # Сохранить на S3 если ещё нет
            if s3 and BUCKET_NAME and not item.get('storage_path'):
                try:
                    digest = hashlib.sha256(bytes(data)).hexdigest()[:16]
                    safe_fn = filename.replace("/", "_").replace("\\", "_")
                    s3_key = f"tg_attachments/{item['table']}/{item['message_id']}/{digest}_{safe_fn}"
                    s3.put_object(Bucket=BUCKET_NAME, Key=s3_key, Body=bytes(data))
                    item['new_storage_path'] = f"s3://{BUCKET_NAME}/{s3_key}"
                except Exception as e:
                    logger.warning(f"S3 upload after TG download failed: {e}")
            
            return bytes(data), filename, 'telegram'
        except Exception as e:
            logger.warning(f"Telegram download failed: {e}")
    
    return None, "", "failed"


# ============================================================
# ФУНКЦИИ АНАЛИЗА
# ============================================================

def analyze_file(gpt_client, file_data, filename, media_type=""):
    """Универсальный анализ файла. Возвращает (analysis, content)."""
    filename_lower = filename.lower()
    
    if media_type in ('photo', 'image') or filename_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp')):
        return analyze_image(gpt_client, file_data, filename)
    elif media_type == 'pdf' or filename_lower.endswith('.pdf'):
        return analyze_pdf(gpt_client, file_data, filename)
    elif filename_lower.endswith(('.xlsx', '.xls')):
        return analyze_excel(gpt_client, file_data, filename)
    elif filename_lower.endswith(('.docx', '.doc')):
        return analyze_word(gpt_client, file_data, filename)
    
    return "", ""


def analyze_image(client, image_data, filename=""):
    analysis = ""
    content = ""
    try:
        mime = 'image/png' if filename.lower().endswith('.png') else 'image/jpeg'
        b64 = base64.standard_b64encode(image_data).decode("utf-8")
        response = client.chat.completions.create(
            model="openai/gpt-4.1",
            max_tokens=1500,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    {"type": "text", "text": f"Проанализируй изображение ({filename}). Опиши содержание, извлеки текст если есть."}
                ]
            }]
        )
        analysis = response.choices[0].message.content
    except Exception as e:
        logger.error(f"analyze_image: {e}")
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(io.BytesIO(image_data))
        content = pytesseract.image_to_string(img, lang='rus+eng').strip()[:5000]
    except Exception:
        pass
    return analysis, content


def analyze_pdf(client, pdf_data, filename=""):
    analysis = ""
    content = ""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))
        parts = []
        for page in reader.pages[:50]:
            t = page.extract_text()
            if t:
                parts.append(t)
        content = '\n'.join(parts)[:20000]
    except Exception:
        pass
    try:
        text_for_llm = content[:8000] if content.strip() else ""
        if len(text_for_llm.strip()) > 50:
            response = client.chat.completions.create(
                model="openai/gpt-4.1",
                max_tokens=4500,
                messages=[{"role": "user", "content": f"Проанализируй PDF ({filename}):\n\n{text_for_llm}"}]
            )
            analysis = response.choices[0].message.content
        else:
            b64 = base64.standard_b64encode(pdf_data).decode("utf-8")
            response = client.chat.completions.create(
                model="openai/gpt-4.1",
                max_tokens=4500,
                messages=[{"role": "user", "content": [
                    {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": b64}},
                    {"type": "text", "text": f"Проанализируй PDF ({filename})."}
                ]}]
            )
            analysis = response.choices[0].message.content
    except Exception as e:
        logger.error(f"analyze_pdf: {e}")
    return analysis, content


def analyze_excel(client, file_data, filename=""):
    analysis = ""
    content = ""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_data), data_only=True)
        rows = []
        for sn in wb.sheetnames[:5]:
            ws = wb[sn]
            sr = []
            for row in ws.iter_rows(max_row=100, values_only=True):
                vals = [str(c) if c is not None else '' for c in row]
                if any(v.strip() for v in vals):
                    sr.append('\t'.join(vals))
            if sr:
                rows.append(f"=== {sn} ===\n" + '\n'.join(sr))
        content = '\n\n'.join(rows)[:20000]
    except Exception:
        pass
    if content.strip():
        try:
            response = client.chat.completions.create(
                model="openai/gpt-4.1", max_tokens=4500,
                messages=[{"role": "user", "content": f"Проанализируй Excel ({filename}):\n\n{content[:8000]}"}]
            )
            analysis = response.choices[0].message.content
        except Exception as e:
            logger.error(f"analyze_excel: {e}")
    return analysis, content


def analyze_word(client, file_data, filename=""):
    analysis = ""
    content = ""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_data))
        content = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())[:20000]
    except Exception:
        pass
    if content.strip():
        try:
            response = client.chat.completions.create(
                model="openai/gpt-4.1", max_tokens=4500,
                messages=[{"role": "user", "content": f"Проанализируй Word ({filename}):\n\n{content[:8000]}"}]
            )
            analysis = response.choices[0].message.content
        except Exception as e:
            logger.error(f"analyze_word: {e}")
    return analysis, content


# ============================================================
# 1. АУДИТ TG-ВЛОЖЕНИЙ
# ============================================================

def find_tg_missing(conn):
    cur = conn.cursor()
    cur.execute("""
        SELECT table_name FROM information_schema.tables 
        WHERE table_schema = 'public' AND table_name LIKE 'tg_chat_%%'
          AND table_name != 'tg_chats_metadata'
    """)
    tables = [r[0] for r in cur.fetchall()]
    
    missing = []
    for tbl in tables:
        if tbl in EXCLUDED_TABLES:
            continue
        try:
            cur.execute(f"""
                SELECT message_id, media_file_id, message_type, timestamp,
                       media_analysis, content_text, storage_path
                FROM {tbl}
                WHERE media_file_id IS NOT NULL AND media_file_id != ''
                  AND (media_analysis IS NULL OR media_analysis = ''
                       OR content_text IS NULL OR content_text = '')
                ORDER BY timestamp DESC
            """)
            for row in cur.fetchall():
                missing.append({
                    'table': tbl,
                    'message_id': row[0],
                    'file_id': row[1],
                    'media_type': row[2] or '',
                    'timestamp': row[3],
                    'has_analysis': bool(row[4] and row[4].strip()),
                    'has_content': bool(row[5] and row[5].strip()),
                    'storage_path': row[6] or '',
                })
        except Exception as e:
            logger.warning(f"Ошибка {tbl}: {e}")
            conn.rollback()
    cur.close()
    return missing


async def fix_tg_media(conn, items):
    from telegram import Bot
    from telegram.request import HTTPXRequest
    
    from proxy_config import get_proxy_url
    request = HTTPXRequest(read_timeout=120, connect_timeout=30, proxy=get_proxy_url())
    bot = Bot(token=BOT_TOKEN, request=request)
    s3 = get_s3_client()
    gpt_client = get_gpt_client()
    
    fixed = 0
    errors = 0
    
    for item in items[:MAX_MEDIA_FIX]:
        try:
            file_data, filename, source = await download_file(item, s3, bot)
            if not file_data:
                errors += 1
                # Помечаем недоступные файлы чтобы не пытаться снова
                try:
                    cur = conn.cursor()
                    updates = []
                    if not item['has_analysis']:
                        updates.append("media_analysis = '[файл недоступен]'")
                    if not item['has_content']:
                        updates.append("content_text = '[файл недоступен]'")
                    if updates:
                        cur.execute(
                            f"UPDATE {item['table']} SET {', '.join(updates)} WHERE message_id = %s",
                            (item['message_id'],)
                        )
                        conn.commit()
                    cur.close()
                except Exception as mark_err:
                    logger.warning(f"  Не удалось пометить недоступный файл {item['table']} msg={item['message_id']}: {mark_err}")
                    conn.rollback()
                continue
            
            logger.info(f"  [{source}] {item['table']} msg={item['message_id']} ({filename})")
            
            analysis, content = analyze_file(gpt_client, file_data, filename, item['media_type'])
            
            updates = []
            params = []
            if analysis and not item['has_analysis']:
                updates.append("media_analysis = %s")
                params.append(analysis)
            if content and not item['has_content']:
                updates.append("content_text = %s")
                params.append(content)
            if item.get('new_storage_path'):
                updates.append("storage_path = %s")
                params.append(item['new_storage_path'])
            
            if updates:
                params.append(item['message_id'])
                cur = conn.cursor()
                cur.execute(f"UPDATE {item['table']} SET {', '.join(updates)} WHERE message_id = %s", params)
                conn.commit()
                cur.close()
                fixed += 1
            
            time.sleep(1)
        except Exception as e:
            logger.error(f"  Ошибка {item['table']} msg={item['message_id']}: {e}")
            errors += 1
            conn.rollback()
    
    return fixed, errors


# ============================================================
# 2. АУДИТ EMAIL-ВЛОЖЕНИЙ
# ============================================================

def audit_email_attachments(conn):
    cur = conn.cursor()
    cur.execute("""
        SELECT COUNT(*) FROM email_attachments 
        WHERE analysis_status = 'pending'
    """)
    count = cur.fetchone()[0]
    cur.close()
    return count


def fix_email_attachments():
    try:
        result = subprocess.run(
            [VENV_PYTHON, f'{SCRIPTS_DIR}/analyze_attachments.py', '--batch', '100'],
            capture_output=True, text=True, timeout=1200, cwd=SCRIPTS_DIR
        )
        for line in result.stdout.split('\n'):
            if 'Done' in line or 'обработано' in line.lower() or 'Processed' in line:
                return line.strip()
        return f"exit={result.returncode}"
    except Exception as e:
        return f"error: {e}"


# ============================================================
# 3. АУДИТ SOURCE_DOCUMENTS
# ============================================================

def audit_source_documents(conn):
    cur = conn.cursor()
    # Дистилляция реализована только для telegram_message и email_message.
    # Остальные source_kinds (email_attachment, matrix_message, c1_event,
    # rag_answer, synthesized_1c) намеренно исключены из этого пайплайна.
    cur.execute("""
        SELECT source_kind, COUNT(*)
        FROM source_documents
        WHERE (meta->>'distilled') IS NULL
          AND LENGTH(body_text) >= 25
          AND source_kind IN ('telegram_message', 'email_message')
          AND (source_kind != 'email_message'
               OR meta->>'email_category' IN ('internal', 'external_business'))
        GROUP BY source_kind
    """)
    undistilled = {row[0]: row[1] for row in cur.fetchall()}
    
    cur.execute("""
        SELECT sd.source_kind, COUNT(*)
        FROM source_documents sd
        LEFT JOIN source_chunks sc ON sc.document_id = sd.id
        WHERE sc.id IS NULL
          AND sd.body_text IS NOT NULL AND LENGTH(sd.body_text) >= 25
          AND (sd.meta->>'skip_reason' IS NULL
               OR sd.meta->>'skip_reason' != 'auto_notification')
        GROUP BY sd.source_kind
    """)
    unchunked = {row[0]: row[1] for row in cur.fetchall()}
    cur.close()
    return undistilled, unchunked


def fix_distillation(undistilled):
    results = []
    for sk in ['telegram_message', 'email_message']:
        cnt = undistilled.get(sk, 0)
        if cnt == 0:
            continue
        batches = min(cnt // 5 + 1, MAX_DISTILL_BATCHES)
        try:
            result = subprocess.run(
                [VENV_PYTHON, f'{SCRIPTS_DIR}/distillation.py', sk, str(batches)],
                capture_output=True, text=True, timeout=1800, cwd=SCRIPTS_DIR
            )
            for line in result.stdout.split('\n'):
                if 'ИТОГО' in line:
                    results.append(f"{sk}: {line.strip()}")
        except Exception as e:
            results.append(f"{sk}: error {e}")
    return results


def fix_chunks():
    try:
        result = subprocess.run(
            [VENV_PYTHON, f'{SCRIPTS_DIR}/build_source_chunks.py', '--batch', str(MAX_CHUNKS_DOCS)],
            capture_output=True, text=True, timeout=1200, cwd=SCRIPTS_DIR
        )
        for line in result.stdout.split('\n'):
            if 'ГОТОВО' in line:
                return line.strip()
        return f"exit={result.returncode}"
    except Exception as e:
        return f"error: {e}"


# ============================================================
# ОТЧЁТ В TELEGRAM
# ============================================================

def send_report(text):
    if not ADMIN_USER_ID or not BOT_TOKEN:
        logger.warning("ADMIN_USER_ID/BOT_TOKEN не заданы")
        return
    url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage"
    for i in range(0, len(text), 4000):
        try:
            requests.post(url, json={
                'chat_id': ADMIN_USER_ID,
                'text': text[i:i+4000],
                'parse_mode': 'HTML'
            }, timeout=30, proxies=get_proxy_dict())
        except Exception as e:
            logger.error(f"Отправка отчёта: {e}")


# ============================================================
# MAIN
# ============================================================

async def main():
    lock_file = open('/tmp/audit_pipeline.lock', 'w')
    try:
        fcntl.flock(lock_file, fcntl.LOCK_EX | fcntl.LOCK_NB)
    except IOError:
        print("audit_pipeline уже запущен")
        sys.exit(0)
    
    logger.info("=" * 60)
    logger.info("АУДИТ ПАЙПЛАЙНА — СТАРТ")
    logger.info("=" * 60)
    
    conn = get_conn()
    report = ["<b>🔍 Ночной аудит пайплайна</b>\n"]
    
    # --- 1. TG-вложения ---
    logger.info("1. TG-вложения...")
    tg_missing = find_tg_missing(conn)
    if tg_missing:
        by_table = {}
        for item in tg_missing:
            by_table.setdefault(item['table'], []).append(item)
        report.append(f"<b>📎 TG-вложения без анализа/содержания:</b> {len(tg_missing)}")
        for tbl, items in sorted(by_table.items(), key=lambda x: -len(x[1])):
            short = tbl.replace('tg_chat_', '')
            parts = short.split('_', 1)
            short = parts[1] if len(parts) > 1 else short
            no_a = sum(1 for i in items if not i['has_analysis'])
            no_c = sum(1 for i in items if not i['has_content'])
            report.append(f"  • {short}: {len(items)} (анализ:{no_a}, текст:{no_c})")
        fixed, errs = await fix_tg_media(conn, tg_missing)
        report.append(f"  ✅ Исправлено: {fixed}, ❌ Ошибок: {errs}")
        if len(tg_missing) > MAX_MEDIA_FIX:
            report.append(f"  ⏳ Осталось: {len(tg_missing) - MAX_MEDIA_FIX}")
    else:
        report.append("📎 TG-вложения: всё ок ✅")
    
    # --- 2. Email-вложения ---
    logger.info("2. Email-вложения...")
    email_missing = audit_email_attachments(conn)
    if email_missing > 0:
        report.append(f"\n<b>📧 Email-вложения без анализа:</b> {email_missing}")
        result = fix_email_attachments()
        report.append(f"  ✅ {result}")
    else:
        report.append("\n📧 Email-вложения: всё ок ✅")
    
    # --- 3. Distillation ---
    logger.info("3. Distillation...")
    undistilled, unchunked = audit_source_documents(conn)
    if undistilled:
        total = sum(undistilled.values())
        report.append(f"\n<b>🧠 Без distillation:</b> {total}")
        for sk, cnt in undistilled.items():
            report.append(f"  • {sk}: {cnt}")
        results = fix_distillation(undistilled)
        for r in results:
            report.append(f"  ✅ {r}")
    else:
        report.append("\n🧠 Distillation: всё ок ✅")
    
    # --- 4. Chunks ---
    logger.info("4. Source chunks...")
    if unchunked:
        total = sum(unchunked.values())
        report.append(f"\n<b>📦 Без chunks:</b> {total}")
        result = fix_chunks()
        report.append(f"  ✅ {result}")
    else:
        report.append("\n📦 Chunks: всё ок ✅")
    
    # --- 5. Статистика ---
    logger.info("5. Статистика...")
    cur = conn.cursor()
    stats = {}
    for tbl in ['km_facts', 'km_decisions', 'km_entities', 'km_tasks', 'km_policies', 'source_chunks']:
        cur.execute(f"SELECT COUNT(*) FROM {tbl}")
        stats[tbl] = cur.fetchone()[0]
    cur.close()
    report.append(f"\n<b>📊 База знаний:</b>")
    report.append(f"  Факты: {stats['km_facts']}, Решения: {stats['km_decisions']}")
    report.append(f"  Сущности: {stats['km_entities']}, Задачи: {stats['km_tasks']}")
    report.append(f"  Политики: {stats['km_policies']}, Чанки: {stats['source_chunks']}")
    
    conn.close()
    
    report_text = '\n'.join(report)
    logger.info("Отчёт:\n" + report_text)
    send_report(report_text)
    logger.info("АУДИТ ЗАВЕРШЁН")


if __name__ == '__main__':
    asyncio.run(main())
