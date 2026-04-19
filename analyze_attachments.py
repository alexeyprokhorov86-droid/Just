#!/usr/bin/env python3
"""
analyze_attachments.py — Анализ email-вложений из S3.

Извлекает текст из PDF/XLSX/XLS/DOCX/CSV, анализирует изображения через Vision,
отправляет на LLM-анализ с контекстом (department, тема письма, роль получателя).

Использование:
    python3 analyze_attachments.py                    # обработать все pending
    python3 analyze_attachments.py --batch 100        # обработать 100 файлов
    python3 analyze_attachments.py --type pdf         # только PDF
    python3 analyze_attachments.py --extract-only      # только извлечение текста, без LLM
    python3 analyze_attachments.py --dry-run           # показать что будет обработано
"""

import os
import sys
import io
import csv
import json
import time
import base64
import logging
import argparse
import tempfile
from datetime import datetime
from pathlib import Path

import boto3
import psycopg2
import psycopg2.extras
from botocore.config import Config
from dotenv import load_dotenv

load_dotenv('/home/admin/telegram_logger_bot/.env')

# ── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('/home/admin/telegram_logger_bot/analyze_attachments.log')
    ]
)
logger = logging.getLogger(__name__)

# ── Config ───────────────────────────────────────────────────────────────────
S3_ENDPOINT = os.getenv('ATTACHMENTS_BUCKET_ENDPOINT')
S3_ACCESS_KEY = os.getenv('ATTACHMENTS_BUCKET_ACCESS_KEY')
S3_SECRET_KEY = os.getenv('ATTACHMENTS_BUCKET_SECRET_KEY')
S3_REGION = os.getenv('ATTACHMENTS_BUCKET_REGION')
S3_BUCKET = os.getenv('ATTACHMENTS_BUCKET_NAME')

DB_HOST = os.getenv('DB_HOST', '172.20.0.2')
DB_PORT = 5432
DB_NAME = 'knowledge_base'
DB_USER = 'knowledge'
DB_PASSWORD = os.getenv('DB_PASSWORD')

ROUTERAI_API_KEY = os.getenv('ROUTERAI_API_KEY')
ROUTERAI_BASE_URL = os.getenv('ROUTERAI_BASE_URL', 'https://routerai.ru/api/v1')

# LLM модели

TEXT_MODEL = 'openai/gpt-4.1-mini' # для анализа текста
VISION_MODEL = 'openai/gpt-4.1-mini'  # для анализа картинок

# Лимиты
MAX_TEXT_FOR_LLM = 8000           # макс символов текста для отправки в LLM
MAX_IMAGE_SIZE_MB = 10            # макс размер изображения для vision
BATCH_PAUSE_SEC = 0.5             # пауза между LLM-запросами

# ── S3 Client ────────────────────────────────────────────────────────────────
def get_s3_client():
    return boto3.client('s3',
        endpoint_url=S3_ENDPOINT,
        aws_access_key_id=S3_ACCESS_KEY,
        aws_secret_access_key=S3_SECRET_KEY,
        region_name=S3_REGION,
        config=Config(s3={'addressing_style': 'path'})
    )

# ── DB Connection ────────────────────────────────────────────────────────────
def get_db_conn():
    return psycopg2.connect(
        host=DB_HOST, port=DB_PORT, dbname=DB_NAME,
        user=DB_USER, password=DB_PASSWORD
    )

# ── Text Extraction ─────────────────────────────────────────────────────────

def extract_text_pdf(filepath):
    """Извлечь текст из PDF через pdfplumber."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
                # Пробуем извлечь таблицы
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            text_parts.append(' | '.join(str(c or '') for c in row))
            except Exception as e:
                logger.warning(f"PDF page {i} error: {e}")
    return '\n'.join(text_parts)


def extract_text_xlsx(filepath):
    """Извлечь текст из XLSX через openpyxl."""
    import openpyxl
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    text_parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"=== Лист: {sheet_name} ===")
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else '' for c in row]
            if any(cells):
                text_parts.append(' | '.join(cells))
                row_count += 1
            if row_count > 500:  # лимит строк на лист
                text_parts.append(f"... (ещё строки, обрезано)")
                break
    wb.close()
    return '\n'.join(text_parts)


def extract_text_xls(filepath):
    """Извлечь текст из XLS через xlrd."""
    import xlrd
    wb = xlrd.open_workbook(filepath)
    text_parts = []
    for sheet in wb.sheets():
        text_parts.append(f"=== Лист: {sheet.name} ===")
        for row_idx in range(min(sheet.nrows, 500)):
            cells = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
            if any(cells):
                text_parts.append(' | '.join(cells))
        if sheet.nrows > 500:
            text_parts.append(f"... (ещё {sheet.nrows - 500} строк)")
    return '\n'.join(text_parts)


def extract_text_docx(filepath):
    """Извлечь текст из DOCX через python-docx."""
    from docx import Document
    doc = Document(filepath)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                text_parts.append(' | '.join(cells))
    return '\n'.join(text_parts)


def extract_text_csv(filepath):
    """Извлечь текст из CSV."""
    text_parts = []
    # Пробуем разные кодировки
    for encoding in ['utf-8', 'cp1251', 'latin-1']:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    text_parts.append(' | '.join(row))
                    if i > 500:
                        text_parts.append("... (обрезано)")
                        break
            break
        except UnicodeDecodeError:
            continue
    return '\n'.join(text_parts)


def extract_text(filepath, ext):
    """Маршрутизатор извлечения текста по расширению."""
    try:
        if ext == 'pdf':
            return extract_text_pdf(filepath)
        elif ext == 'xlsx' or ext == 'xlsm':
            return extract_text_xlsx(filepath)
        elif ext == 'xls':
            return extract_text_xls(filepath)
        elif ext in ('docx', 'doc'):
            if ext == 'doc':
                logger.info(f"  .doc format — skipping (no parser)")
                return None
            return extract_text_docx(filepath)
        elif ext == 'csv':
            return extract_text_csv(filepath)
        else:
            return None
    except Exception as e:
        logger.error(f"  Text extraction error ({ext}): {e}")
        return None


# ── LLM Analysis ────────────────────────────────────────────────────────────

def call_llm(messages, model=TEXT_MODEL, max_tokens=1500):
    """Вызов LLM через RouterAI."""
    import urllib.request
    
    payload = json.dumps({
        "model": model,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": 0.3
    }).encode('utf-8')
    
    req = urllib.request.Request(
        f"{ROUTERAI_BASE_URL}/chat/completions",
        data=payload,
        headers={
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {ROUTERAI_API_KEY}'
        }
    )
    
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            return data['choices'][0]['message']['content']
    except Exception as e:
        logger.error(f"  LLM call error: {e}")
        return None


def call_vision(image_base64, media_type, context_prompt):
    """Вызов Vision LLM для анализа изображения."""
    import urllib.request
    
    messages = [
        {"role": "system", "content": "Ты аналитик компании Фрумелад (производство кондитерских изделий). Анализируй вложения из деловой переписки. Отвечай на русском. Будь конкретен и краток."},
        {"role": "user", "content": [
            {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{image_base64}"}},
            {"type": "text", "text": context_prompt}
        ]}
    ]
    
    payload = json.dumps({
        "model": VISION_MODEL,
        "messages": messages,
        "max_tokens": 1000,
        "temperature": 0.3
    }).encode('utf-8')
    
    req = urllib.request.Request(
        f"{ROUTERAI_BASE_URL}/chat/completions",
        data=payload,
        headers={
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {ROUTERAI_API_KEY}'
        }
    )
    
    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            return data['choices'][0]['message']['content']
    except Exception as e:
        logger.error(f"  Vision call error: {e}")
        return None


def build_context_prompt(attachment, extracted_text=None):
    """Формирует промпт с контекстом для LLM-анализа."""
    dept = attachment.get('department', 'неизвестно')
    role_desc = attachment.get('role_description', '')
    employee = attachment.get('employee_name', '')
    subject = attachment.get('subject', '')
    from_addr = attachment.get('from_address', '')
    mailbox = attachment.get('mailbox_email', '')
    filename = attachment.get('filename', '')
    
    context_parts = [
        f"Контекст письма:",
        f"- Тема: {subject}" if subject else None,
        f"- От: {from_addr}" if from_addr else None,
        f"- Получатель (ящик): {mailbox}" if mailbox else None,
        f"- Департамент получателя: {dept}" if dept and dept != 'общее' else None,
        f"- Роль получателя: {role_desc}" if role_desc and 'общего назначения' not in role_desc else None,
        f"- Сотрудник: {employee}" if employee else None,
        f"- Файл: {filename}",
    ]
    context = '\n'.join(p for p in context_parts if p)
    
    if extracted_text:
        text_truncated = extracted_text[:MAX_TEXT_FOR_LLM]
        if len(extracted_text) > MAX_TEXT_FOR_LLM:
            text_truncated += "\n... (текст обрезан)"
        
        return f"""{context}

Содержимое файла:
{text_truncated}

Задача: Кратко опиши что это за документ, его ключевую информацию (суммы, даты, контрагенты, товары, условия). 
Если это счёт/накладная/акт — укажи номер, дату, сумму, контрагентов.
Если это отчёт/таблица — опиши структуру и ключевые данные.
Ответь 3-5 предложениями."""
    else:
        return f"""{context}

Задача: Опиши что изображено на этой картинке в контексте деловой переписки кондитерской компании.
Если это документ/скан — извлеки ключевую информацию.
Если это фото продукции/упаковки/оборудования — опиши что видишь.
Если это скриншот — опиши содержание.
Ответь 2-4 предложениями."""


# ── Main Processing ─────────────────────────────────────────────────────────

def get_pending_attachments(conn, batch_size=None, file_type=None):
    """Получить список pending вложений с контекстом."""
    cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
    
    type_filter = ""
    if file_type:
        ext_map = {
            'pdf': ("'%.pdf'",),
            'xlsx': ("'%.xlsx'", "'%.xlsm'"),
            'xls': ("'%.xls'",),
            'docx': ("'%.docx'",),
            'csv': ("'%.csv'",),
            'image': ("'%.png'", "'%.jpg'", "'%.jpeg'", "'%.heic'"),
        }
        patterns = ext_map.get(file_type, (f"'%.{file_type}'",))
        conditions = ' OR '.join(f"LOWER(ea.filename) LIKE {p}" for p in patterns)
        type_filter = f"AND ({conditions})"
    
    limit_clause = f"LIMIT {batch_size}" if batch_size else ""
    
    cur.execute(f"""
        SELECT 
            ea.id, ea.filename, ea.storage_path, ea.size_bytes, ea.content_type,
            e.subject, e.from_address, e.category,
            mb.email as mailbox_email,
            em.department, em.role_description, em.employee_name, em.is_active as mailbox_active
        FROM email_attachments ea
        JOIN email_messages e ON e.id = ea.message_id
        JOIN monitored_mailboxes mb ON mb.id = e.mailbox_id
        LEFT JOIN email_employee_mapping em ON em.email = mb.email
        WHERE ea.analysis_status = 'pending'
        {type_filter}
        ORDER BY ea.size_bytes ASC
        {limit_clause}
    """)
    
    return cur.fetchall()


def process_attachment(s3_client, conn, att, extract_only=False):
    """Обработать одно вложение."""
    att_id = att['id']
    filename = att['filename']
    s3_path = att['storage_path']
    size = att['size_bytes']
    
    ext = (filename.rsplit('.', 1)[-1] if '.' in filename else '').lower()
    is_image = ext in ('png', 'jpg', 'jpeg', 'heic', 'bmp')
    is_text_extractable = ext in ('pdf', 'xlsx', 'xlsm', 'xls', 'docx', 'csv')
    
    logger.info(f"[{att_id}] {filename} ({size} bytes) ext={ext}")
    
    # Скачиваем из S3
    s3_key = s3_path.replace(f's3://{S3_BUCKET}/', '')
    
    try:
        with tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False) as tmp:
            tmp_path = tmp.name
            s3_client.download_file(S3_BUCKET, s3_key, tmp_path)
    except Exception as e:
        logger.error(f"  S3 download error: {e}")
        update_attachment_status(conn, att_id, 'error', error=str(e))
        return False
    
    content_text = None
    analysis_text = None
    
    try:
        # Этап 1: Извлечение текста
        if is_text_extractable:
            content_text = extract_text(tmp_path, ext)
            if content_text:
                logger.info(f"  Extracted {len(content_text)} chars")
            else:
                logger.info(f"  No text extracted")
        
        # Этап 2: LLM-анализ
        if not extract_only:
            if is_image:
                # Vision для изображений
                if size > MAX_IMAGE_SIZE_MB * 1024 * 1024:
                    logger.info(f"  Image too large for vision ({size} bytes)")
                    update_attachment_status(conn, att_id, 'skip_large', content_text=content_text)
                    return True
                
                with open(tmp_path, 'rb') as f:
                    image_data = base64.b64encode(f.read()).decode('utf-8')
                
                media_type = {
                    'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
                    'heic': 'image/heic', 'bmp': 'image/bmp'
                }.get(ext, 'image/png')
                
                context_prompt = build_context_prompt(att)
                analysis_text = call_vision(image_data, media_type, context_prompt)
                
                if analysis_text:
                    logger.info(f"  Vision analysis: {len(analysis_text)} chars")
                
                time.sleep(BATCH_PAUSE_SEC)
            
            elif content_text and len(content_text.strip()) > 50:
                # LLM-анализ текста
                context_prompt = build_context_prompt(att, content_text)
                messages = [
                    {"role": "system", "content": "Ты аналитик компании Фрумелад (производство кондитерских изделий). Анализируй вложения из деловой переписки. Отвечай на русском. Будь конкретен и краток."},
                    {"role": "user", "content": context_prompt}
                ]
                analysis_text = call_llm(messages)
                
                if analysis_text:
                    logger.info(f"  LLM analysis: {len(analysis_text)} chars")
                
                time.sleep(BATCH_PAUSE_SEC)
        
        # Сохраняем результат
        status = 'done' if (content_text or analysis_text) else 'empty'
        update_attachment_status(
            conn, att_id, status,
            content_text=content_text,
            analysis_text=analysis_text,
            model=VISION_MODEL if is_image else TEXT_MODEL
        )
        return True
        
    except Exception as e:
        logger.error(f"  Processing error: {e}")
        update_attachment_status(conn, att_id, 'error', error=str(e))
        return False
    finally:
        # Удаляем tmp файл
        try:
            os.unlink(tmp_path)
        except:
            pass


def update_attachment_status(conn, att_id, status, content_text=None, analysis_text=None, model=None, error=None):
    """Обновить статус вложения в БД + canonicalize при успехе."""
    cur = conn.cursor()
    cur.execute("""
        UPDATE email_attachments SET
            analysis_status = %s,
            content_text = COALESCE(%s, content_text),
            analysis_text = COALESCE(%s, analysis_text),
            analysis_model = COALESCE(%s, analysis_model),
            analyzed_at = NOW(),
            analysis_error = %s
        WHERE id = %s
    """, (status, content_text, analysis_text, model, error, att_id))
    conn.commit()
    if status == 'done':
        try:
            insert_email_attachment_to_canonical(conn, att_id)
        except Exception as e:
            logger.warning(f"  canonical insert failed for att_id={att_id}: {e}")


def insert_email_attachment_to_canonical(conn, att_id: int):
    """INSERT email_attachment в source_documents (idempotent)."""
    import json
    cur = conn.cursor()
    cur.execute("""
        SELECT a.id, a.message_id, a.filename, a.content_type, a.size_bytes,
               a.storage_path, a.analysis_text, a.content_text, a.analysis_model,
               em.received_at, em.from_address,
               mb.email AS mailbox_email,
               sd.id AS parent_source_doc_id
        FROM email_attachments a
        JOIN email_messages em ON em.id = a.message_id
        LEFT JOIN monitored_mailboxes mb ON mb.id = em.mailbox_id
        LEFT JOIN source_documents sd ON sd.source_kind='email_message'
                                     AND sd.source_ref = 'email:' || em.id
        WHERE a.id = %s
    """, (att_id,))
    row = cur.fetchone()
    if not row:
        return
    (aid, em_id, filename, ctype, size_b, storage, analysis_t, content_t,
     model, received_at, from_addr, mailbox_email, parent_sd_id) = row

    parts = []
    if analysis_t:
        parts.append(f"[Анализ файла: {filename}]\n{analysis_t}")
    if content_t:
        parts.append(f"---ПОЛНЫЙ ТЕКСТ---\n\n{content_t}")
    body = '\n\n'.join(parts).strip()
    if not body or len(body) < 25:
        return

    if ctype and ctype.startswith('image/'): extr = 'vision'
    elif ctype == 'application/pdf': extr = 'pdf_text'
    else: extr = 'extract'
    meta = {
        'parent_email_id': em_id,
        'parent_source_doc_id': parent_sd_id,
        'content_type': ctype or '',
        'size_bytes': size_b,
        'storage_path': storage,
        'analysis_model': model,
        'extraction_method': extr,
        'has_extracted_text': bool(content_t),
        'has_llm_summary': bool(analysis_t),
    }
    cur.execute("""
        INSERT INTO source_documents
            (source_kind, source_ref, title, body_text, doc_date,
             author_name, author_ref, channel_ref, channel_name,
             language, is_deleted, confidence, meta)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        ON CONFLICT (source_kind, source_ref) DO NOTHING
    """, (
        'email_attachment', f'email_att:{aid}', filename, body, received_at,
        from_addr, from_addr, mailbox_email, mailbox_email,
        'ru', False, 1.0, json.dumps(meta)
    ))
    conn.commit()


# ── CLI ──────────────────────────────────────────────────────────────────────

def main():
    # Защита от одновременного запуска
    import fcntl
    lock_file = open('/tmp/analyze_attachments.lock', 'w')
    try:
        fcntl.flock(lock_file, fcntl.LOCK_EX | fcntl.LOCK_NB)
    except IOError:
        logger.info("Already running, exiting")
        return
    
    parser = argparse.ArgumentParser(description='Analyze email attachments from S3')
    parser.add_argument('--batch', type=int, default=None, help='Number of files to process')
    parser.add_argument('--type', type=str, default=None, 
                       choices=['pdf', 'xlsx', 'xls', 'docx', 'csv', 'image'],
                       help='Process only this file type')
    parser.add_argument('--extract-only', action='store_true', help='Only extract text, no LLM')
    parser.add_argument('--dry-run', action='store_true', help='Show what would be processed')
    args = parser.parse_args()
    
    conn = get_db_conn()
    attachments = get_pending_attachments(conn, batch_size=args.batch, file_type=args.type)
    
    logger.info(f"Found {len(attachments)} pending attachments" + 
                (f" (type={args.type})" if args.type else ""))
    
    if args.dry_run:
        for att in attachments[:20]:
            print(f"  [{att['id']}] {att['filename']} ({att['size_bytes']} bytes) "
                  f"dept={att['department']} from={att['from_address']}")
        if len(attachments) > 20:
            print(f"  ... and {len(attachments) - 20} more")
        conn.close()
        return
    
    if not attachments:
        logger.info("Nothing to process")
        conn.close()
        return
    
    s3 = get_s3_client()
    
    success = 0
    errors = 0
    start_time = time.time()
    
    for i, att in enumerate(attachments):
        logger.info(f"--- [{i+1}/{len(attachments)}] ---")
        try:
            ok = process_attachment(s3, conn, att, extract_only=args.extract_only)
            if ok:
                success += 1
            else:
                errors += 1
        except Exception as e:
            logger.error(f"Unexpected error: {e}")
            errors += 1
        
        # Прогресс каждые 50 файлов
        if (i + 1) % 50 == 0:
            elapsed = time.time() - start_time
            rate = (i + 1) / elapsed * 60
            logger.info(f"Progress: {i+1}/{len(attachments)} "
                       f"({success} ok, {errors} err) "
                       f"rate: {rate:.0f}/min")
    
    elapsed = time.time() - start_time
    logger.info(f"Done: {success} ok, {errors} errors in {elapsed:.0f}s")
    conn.close()


if __name__ == '__main__':
    main()
